import subprocess
import json
import fitz  # PyMuPDF
from docx import Document

def read_pdf_and_summarize(pdf_path, output_docx_path, start_page=1):
    """
    Reads a PDF file and creates a Word document with summaries of each section.
    
    Args:
        pdf_path: Path to input PDF file
        output_docx_path: Path where output DOCX will be saved 
        start_page: Page number to start processing from (default=1)
    """
    pdf_doc = fitz.open(pdf_path)
    word_doc = Document()
    
    start_page = max(0, start_page) - 1
    
    for page_num in range(start_page, len(pdf_doc)):
        page = pdf_doc.load_page(page_num)
        print(f"Processing page {page_num + 1}...")
        blocks = page.get_text("dict")["blocks"]
        section_text = ""
        section_title = ""
        
        for block in blocks:
            if block["type"] == 0:  # Text block
                block_text = ""
                max_font_size = 0

                for line in block["lines"]:
                    for span in line["spans"]:
                        block_text += span["text"] + " "
                        max_font_size = max(max_font_size, span["size"])
                
                block_text = block_text.strip()

                # Assume text with font size > 13 is a heading
                if max_font_size > 13:
                    if section_text.strip():
                        # Only summarize sections with enough content
                        words = section_text.split()
                        if len(words) >= 35:
                            summary = send_to_api_for_summary(section_text)
                            word_doc.add_heading(section_title, level=1)
                            word_doc.add_paragraph(summary)

                    # Handle new title    
                    if section_text.strip():
                        section_title = block_text
                        section_text = ""
                    else:
                        section_title += block_text

                else:
                    section_text += " " + block_text
        
        # Process final section of the page
        if section_text.strip():
            summary = send_to_api_for_summary(section_text)
            word_doc.add_heading(section_title, level=1)
            word_doc.add_paragraph(summary)
        
        word_doc.save(output_docx_path)
    
    pdf_doc.close()
    print(f"Word document saved to: {output_docx_path}")

def send_to_api_for_summary(text):
    """
    Splits text into chunks and sends to API for summarization
    """
    words = text.split()
    for i in range(0, len(words), 1900):
        chunk = " ".join(words[i:i+1900])
        summary = get_summary_from_api(chunk)
        return summary

def escape_text_for_curl(text):
    """
    Escapes special characters that could cause issues in shell commands
    """
    chars_to_escape = {
        '"': '\\"',  # Escapa comillas dobles
        "'": "\\'",  # Escapa comillas simples
        "&": "and",  # Escapa ampersand
        "`": "",  # Escapa backticks
        "$": "dollars",  # Escapa símbolo de dólar
        ">": "(morethan)",  # Escapa símbolo de dólar
        "<": "(lessthan)",  # Escapa símbolo de dólar
        "": "",  # Escapa símbolo de dólar
        "  ": "",  # Escapa símbolo de dólar
        "": "x",  # Escapa símbolo de dólar
        "/": "-divide-",  # Escapa backslashes
        "“": "",  # Escapa backslashes
        "”": "",  # Escapa backslashes
        "➔": ":-",  # Escapa backslashes
        "“": "",  # Escapa backslashes
        "": "",
    }
    for char, escaped in chars_to_escape.items():
        text = text.replace(char, escaped)
    return text

def get_summary_from_api(text):
    """
    Sends text to local LLM API for processing
    Returns the API response or error message
    """
    escaped_text = escape_text_for_curl(text)
    
    # API call to translate text to English
    curl_cmd = f'curl -X POST http://localhost:1234/v1/chat/completions -H "Content-Type: application/json" -d "{{\\"messages\\": [{{\\"role\\": \\"system\\", \\"content\\": \\"Translate this text to english, be direct, I don\'t want sidenotes: \\"}},{{\\"role\\": \\"user\\", \\"content\\": \\"{escaped_text}\\"}}], \\"temperature\\": 0.7, \\"max_tokens\\": -1, \\"stream\\": false}}"'

    try:
        process = subprocess.run(curl_cmd, shell=True, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, encoding='utf-8')
        result = process.stdout

        if not result:
            print("Empty response or error occurred.")
            return "Empty response or error occurred."

        response_json = json.loads(result)
        message = response_json['choices'][0]['message']['content']
        return message

    except json.JSONDecodeError as e:
        print(f"JSON decode error: {e}")
        return "Error: Failed to decode JSON response."
    except KeyError as e:
        print(f"JSON structure error: {e}")
        return "Error: Unexpected JSON structure."
    except Exception as e:
        print(f"Unexpected error: {e}")
        return str(e)

# Example usage
pdf_path = "document.pdf" 
summaries = read_pdf_and_summarize(pdf_path, pdf_path[:-4] + "_summary.docx")
